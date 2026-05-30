"""Create a Word document from a natural-language goal.

The primary path asks the configured MiniMax model for final document content.
The fallback path still creates a usable document, but the app should prefer AI
content whenever a configured API key is available.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from platformdirs import user_config_dir

try:
    import tomllib
except ModuleNotFoundError:  # pragma: no cover - Python 3.11+ in normal runtime
    import tomli as tomllib  # type: ignore[no-redef]


APP_NAME = "iAgent"
DEFAULT_MODEL = "MiniMax-M2.7"


@dataclass
class Section:
    heading: str
    paragraphs: list[str]


@dataclass
class GeneratedDocument:
    title: str
    sections: list[Section]


def normalize_goal(goal: str) -> str:
    clean = goal.strip()
    clean = re.sub(r"(?i)\bwith\s+proper\s+formatting\b", "", clean)
    clean = re.sub(
        r"(?i)\bsave\s+as\s+[\"'][^\"']+[\"']\s*(?:in\s+the\s+user'?s\s+documents\s+folder)?",
        "",
        clean,
    )
    clean = re.sub(
        r"(?i)\bsave\s+as\s+[A-Za-z0-9_.\-]+\s*(?:in\s+the\s+user'?s\s+documents\s+folder)?",
        "",
        clean,
    )
    clean = re.sub(r"(?i)\b(and|then)\s+open\s+(?:it|the\s+result|in\s+word)\b", "", clean)
    clean = re.sub(r"\s+", " ", clean).strip(" .,;:")
    return clean or goal.strip()


def page_count_from_goal(goal: str) -> int:
    match = re.search(r"(\d+)\s*-?\s*(?:pages?|page\s+long)\b", goal, re.IGNORECASE)
    if not match:
        return 3
    return max(1, min(25, int(match.group(1))))


def subject_from_goal(goal: str) -> str:
    titled = re.search(r'titled\s+"([^"]+)"', goal, re.IGNORECASE)
    if titled:
        return titled.group(1).strip()
    match = re.search(
        r"(?:about|on|regarding|concerning)\s+(.+?)(?=(?:\s+(?:and|with|including|covering|save|open|then|please)\b)|[.;:\n\r]|$)",
        goal,
        re.IGNORECASE,
    )
    if match and match.group(1).strip():
        return match.group(1).strip()
    cleaned = re.sub(r"(?i)\b(write|create|make|draft|generate|prepare)\b", "", goal)
    cleaned = re.sub(r"(?i)\b(a|an|the|\d+\s*-?\s*pages?|word|document|docx|essay|report)\b", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .,;:")
    return cleaned or "Document"


def title_from_goal(goal: str) -> str:
    titled = re.search(r'titled\s+"([^"]+)"', goal, re.IGNORECASE)
    if titled:
        return titled.group(1).strip()
    subject = subject_from_goal(goal)
    if subject and subject.lower() != "document":
        return subject[:1].upper() + subject[1:]
    return "Document"


def safe_filename(title: str) -> str:
    clean = re.sub(r'[<>:"/\\|?*]', "", title).strip()
    clean = re.sub(r"\s+", " ", clean)
    return (clean or "Document")[:80].strip()


def output_dir(default: str | None) -> Path:
    if default:
        return Path(default).expanduser()
    return Path.home() / "Documents" / "iAgent Documents"


def config_path() -> Path:
    return Path(user_config_dir(APP_NAME, appauthor=False, roaming=True)) / "config.toml"


def minimax_key_from_config() -> str | None:
    env_key = os.environ.get("MINIMAX_API_KEY", "").strip()
    if env_key:
        return env_key
    path = config_path()
    if not path.is_file():
        return None
    try:
        data = tomllib.loads(path.read_text(encoding="utf-8"))
    except (OSError, tomllib.TOMLDecodeError, UnicodeDecodeError):
        return None
    key = str(data.get("minimax_api_key", "")).strip()
    return key or None


def target_word_count(page_count: int) -> int:
    return max(1200, min(22000, page_count * 850))


def generation_prompt(goal: str, page_count: int) -> str:
    words = target_word_count(page_count)
    sections = max(5, min(20, page_count * 2))
    min_paragraphs_per_section = max(4, page_count)
    return f"""Create a comprehensive, full-length Microsoft Word document.

User request:
{goal}

Requirements:
- Write only the final polished document content. Do not include reasoning, self-correction, or instructions to the user.
- Match the subject, tone, and format the user requested.
- Do not mention file paths, saving, opening, formatting instructions, or that the document was generated.
- The target is approximately {words} words for about {page_count} page(s). Produce a full, substantive document — not a summary or outline.
- Structure the document with approximately {sections} sections.
- Each section must contain at least {min_paragraphs_per_section} substantial paragraphs of several sentences each.
- Each paragraph should be 80-150 words and must contain real content — no bullet points, no placeholders, no vague statements.
- Include specific details, concrete examples, nuanced analysis, and smooth transitions between paragraphs.
- Vary paragraph length and content to avoid repetition and maintain reader engagement.
- Use exactly this JSON structure — no markdown code fences, no additional fields:
{{
  "title": "Document title",
  "sections": [
    {{"heading": "Section heading", "paragraphs": ["paragraph one (80-150 words, substantive)", "paragraph two (80-150 words)", "paragraph three (80-150 words)", "paragraph four (80-150 words)"]}}
  ]
}}
- Every paragraph must be independently meaningful and developed. Do not leave any paragraph as a placeholder or stub.
""".strip()


def call_minimax(goal: str, page_count: int, api_key: str) -> str:
    payload = {
        "model": DEFAULT_MODEL,
        "max_tokens": max(4000, min(32000, page_count * 2800)),
        "stream": False,
        "reasoning_split": True,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a careful long-form document writer. "
                    "Return polished final document content only."
                ),
            },
            {"role": "user", "content": generation_prompt(goal, page_count)},
        ],
    }
    response = httpx.post(
        "https://api.minimax.io/v1/chat/completions",
        json=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        timeout=180.0,
        follow_redirects=True,
    )
    response.raise_for_status()
    data = response.json()
    return str(data["choices"][0]["message"]["content"])


def extract_json_object(text: str) -> dict | None:
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        value = json.loads(cleaned)
        return value if isinstance(value, dict) else None
    except json.JSONDecodeError:
        pass

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        value = json.loads(cleaned[start : end + 1])
    except json.JSONDecodeError:
        return None
    return value if isinstance(value, dict) else None


def document_from_json(data: dict, fallback_title: str) -> GeneratedDocument | None:
    title = str(data.get("title") or fallback_title).strip() or fallback_title
    raw_sections = data.get("sections")
    if not isinstance(raw_sections, list):
        return None
    sections: list[Section] = []
    for raw in raw_sections:
        if not isinstance(raw, dict):
            continue
        heading = str(raw.get("heading") or "").strip()
        paragraphs_raw = raw.get("paragraphs")
        if not heading or not isinstance(paragraphs_raw, list):
            continue
        paragraphs = [str(p).strip() for p in paragraphs_raw if str(p).strip()]
        if paragraphs:
            sections.append(Section(heading=heading, paragraphs=paragraphs))
    if not sections:
        return None
    return GeneratedDocument(title=title, sections=sections)


def document_from_plain_text(text: str, fallback_title: str) -> GeneratedDocument:
    lines = [line.strip() for line in text.splitlines()]
    title = fallback_title
    sections: list[Section] = []
    current_heading = ""
    current_paragraphs: list[str] = []
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        if buffer:
            current_paragraphs.append(" ".join(buffer).strip())
            buffer = []

    def flush_section() -> None:
        nonlocal current_heading, current_paragraphs
        flush_paragraph()
        if current_heading and current_paragraphs:
            sections.append(Section(current_heading, current_paragraphs))
        current_heading = ""
        current_paragraphs = []

    for line in lines:
        if not line:
            flush_paragraph()
            continue
        if line.startswith("# "):
            title = line[2:].strip() or title
            continue
        if line.startswith("## "):
            flush_section()
            current_heading = line[3:].strip()
            continue
        if not current_heading:
            current_heading = "Overview"
        buffer.append(line)
    flush_section()

    if not sections:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        sections = [Section("Overview", paragraphs or [text.strip()])]
    return GeneratedDocument(title=title, sections=sections)


def detect_style(goal: str) -> str:
    low = goal.lower()
    if re.search(r"\b(love|heart|emotion|relationship|philosophy|essay|poem|art|literature)\b", low):
        return "humanities"
    if re.search(r"\b(api|software|architecture|engineering|technical|implementation|system|model|agi|ai)\b", low):
        return "technical"
    return "general"


def fallback_document(goal: str, page_count: int) -> GeneratedDocument:
    subject = subject_from_goal(goal)
    title = title_from_goal(goal)
    style = detect_style(goal)
    subject_display = subject if subject[:1].isupper() or subject.isupper() else subject.capitalize()
    subject_l = subject[:1].lower() + subject[1:] if subject else "the topic"
    if style == "humanities":
        headings = [
            f"Introduction: Why {title} Matters",
            "The Language of Attachment",
            "Emotional Depth and Vulnerability",
            "Philosophical Perspectives",
            "Love in Daily Life",
            "Conflict, Patience, and Repair",
            "Family, Friendship, and Community",
            "Culture, Art, and Memory",
            "The Ethics of Care",
            "Conclusion: A Practice of Attention",
        ]
        paragraph_sets = [
            [
                "{subject} is not only an emotion that arrives suddenly; it is also a way of paying attention. It changes what a person notices, what they remember, and what they are willing to protect. When people speak about {subject_l}, they often describe warmth or attraction, but underneath those feelings is a deeper recognition: another life has become meaningful enough to alter one's own.",
                "This makes {subject_l} difficult to reduce to a single definition. It can be tenderness, desire, loyalty, forgiveness, devotion, or grief. It can be quiet and steady, or urgent and overwhelming. A serious essay on {subject_l} has to hold these tensions together instead of pretending that one version explains them all.",
                "The importance of {subject_l} is visible in ordinary choices. People make time, change habits, listen more carefully, and sometimes sacrifice comfort because another person matters. The feeling becomes real through behavior, and behavior gives the feeling a durable shape.",
                "For that reason, {subject_l} is both personal and practical. It belongs to the heart, but it also appears in kitchens, hospital rooms, apologies, shared plans, and small acts of patience. Its beauty comes from emotion, yet its meaning is proven in action.",
            ],
            [
                "Attachment gives {subject_l} a structure. Human beings are not built to live as sealed, self-sufficient units; they seek safety, recognition, and response. Love grows where someone feels seen without being reduced, known without being controlled, and accepted without being ignored.",
                "This attachment can be joyful, but it can also be frightening. To love is to admit that another person's absence, pain, or disappointment can affect one's own inner life. Vulnerability is therefore not a weakness beside love; it is one of love's conditions.",
                "Healthy attachment does not erase individuality. The strongest bonds often give people more room to become themselves, because trust reduces the need for constant self-defense. In this sense, love is not possession. It is a relationship in which closeness and freedom learn to coexist.",
                "The language of attachment also explains why love can hurt. A broken promise, a silence, or a loss matters because the bond has become part of one's emotional world. Love deepens life by opening it, and anything open can be wounded.",
            ],
            [
                "Philosophically, {subject_l} raises questions about what makes a life good. Many traditions describe love as a force that pulls people beyond selfishness, teaching them to value something outside their own immediate desire. It asks the self to expand rather than disappear.",
                "At the same time, love is not pure self-sacrifice. A relationship that requires one person to vanish is not love at its healthiest; it is imbalance dressed in noble language. Mature love includes care for the other and respect for the self, because both people must remain real.",
                "This balance is why love has often been connected to truth. To love someone well is not to idealize them into perfection, but to see them with generosity and clarity. It means recognizing flaws without turning the whole person into a flaw.",
                "Love therefore becomes a practice of judgment. It decides when to forgive, when to wait, when to speak honestly, and when to let go. These decisions are rarely simple, which is why love remains one of the central subjects of philosophy, literature, and memory.",
            ],
            [
                "In daily life, {subject_l} is often less dramatic than people imagine. It appears as reliability: arriving when promised, remembering what matters, doing the repetitive work that care requires. These actions may not look poetic, but they are the grammar from which lasting affection is built.",
                "Love also changes the experience of time. Shared routines become meaningful because they are repeated with someone who matters. A meal, a walk, or a conversation can carry emotional weight not because it is rare, but because it is familiar and chosen again.",
                "Practical love includes communication. People who care for each other still misunderstand each other, and affection alone does not solve confusion. Listening, naming needs, apologizing, and repairing trust are skills that allow emotion to survive contact with real life.",
                "This practical dimension does not make love less romantic. It makes romance more believable. Grand feeling is easier to admire than daily patience, but the latter is often where love proves its depth.",
            ],
            [
                "Conflict is not automatically the opposite of {subject_l}. In many relationships, conflict reveals where needs, fears, and expectations have been left unspoken. What matters is not whether disagreement appears, but whether people can remain humane while facing it.",
                "Repair is one of love's most important arts. A sincere apology does more than admit a mistake; it says the relationship is worth the discomfort of honesty. Forgiveness, when it is healthy, does not erase harm. It creates a path by which trust can be rebuilt through changed behavior.",
                "Patience matters because people do not grow at identical speeds. Love often asks one person to wait without becoming passive, and another to change without being shamed into silence. The difficulty is finding compassion that does not abandon responsibility.",
                "Some relationships, however, cannot be repaired safely. Love should not require enduring cruelty or neglect. The ethical side of love includes knowing when care must include distance, boundaries, or an ending.",
            ],
            [
                "Love is larger than romance. Family bonds, friendships, mentorship, and community care all show that love can organize human life in many forms. Each form has its own rituals, promises, and expectations, but all involve the recognition that another person's well-being matters.",
                "Friendship is especially important because it reveals love without the script of romance. Friends choose each other through conversation, shared history, humor, loyalty, and presence. Their love may be less publicly celebrated, but it can be one of the most stable forms of human belonging.",
                "Communal love appears when people care for those beyond their immediate circle. It is visible in mutual aid, hospitality, teaching, nursing, and the defense of dignity. This kind of love turns private tenderness into public responsibility.",
                "Seeing these different forms prevents a narrow view of love. It reminds us that the human need to give and receive care is not confined to one relationship model. Love is a network of practices that helps people survive, grow, and make meaning together.",
            ],
            [
                "Art and culture preserve love because love is one of the experiences people most need to understand after it happens. Songs, novels, paintings, films, and letters become containers for feelings that ordinary speech cannot always carry.",
                "These expressions also shape expectations. Culture can teach people to imagine love as rescue, destiny, partnership, sacrifice, or liberation. Some of these stories are beautiful; others are dangerous when they turn control or obsession into romance.",
                "Memory gives love another dimension. People often continue to be shaped by those they have loved even after distance or death. A phrase, a habit, a place, or an object can keep emotional connection alive in subtle ways.",
                "The cultural life of love shows that it is not merely private chemistry. It is a shared human language, constantly revised by each generation as people try to name what they long for and what they have lost.",
            ],
            [
                "The ethics of {subject_l} begin with care. To love someone is to take their reality seriously: their limits, fears, hopes, and autonomy. Care without respect can become control, while respect without care can become distance. Love asks for both.",
                "This ethical demand is why love cannot be measured only by intensity. A feeling can be powerful and still be selfish. The question is whether it helps another person flourish, whether it allows truth, and whether it accepts responsibility for its effects.",
                "Love also teaches humility. No one fully knows another person, and no relationship is immune from misunderstanding. Humility keeps love curious, preventing it from hardening into assumption or entitlement.",
                "At its best, {subject_l} becomes a discipline of attention. It trains people to notice what another life needs while remaining honest about their own. That is why love is not only something people fall into; it is something they learn to practice.",
            ],
        ]
        expansion = [
            "Viewed through {heading_l}, {subject_l} also shows how deeply people need continuity. A relationship becomes meaningful because it creates a thread between past, present, and future: remembered kindness, present trust, and the hope that care will continue. This continuity helps explain why love can feel like home even when life around it remains uncertain.",
            "There is also a moral seriousness in {subject_l}. Love asks people to notice consequences that selfishness would prefer to ignore. Words, silences, promises, and absences all matter more when another person is emotionally involved. In that sense, love educates attention by making indifference harder to defend.",
            "The practical challenge is that love has to live inside imperfect conditions. People are tired, distracted, proud, afraid, and shaped by histories they do not always understand. A mature view of {subject_l} does not deny these limits. It asks how care can remain honest and generous despite them.",
            "This is why {subject_l} remains such a lasting subject for reflection. It is intimate enough to feel private, but universal enough to connect almost every human story. Whether it appears as romance, friendship, family devotion, or compassion, it reveals what people value when they are most open.",
        ]
    elif style == "technical":
        headings = [
            "Overview",
            "Core Concepts",
            "Current State",
            "Opportunities",
            "Risks and Constraints",
            "Implementation Considerations",
            "Future Outlook",
            "Conclusion",
        ]
        paragraph_sets = [
            [
                "{subject} should be understood through the concrete capabilities it enables and the limits that still shape its use. A strong overview starts by separating aspiration from implementation: what the system is expected to do, where the evidence is strongest, and where uncertainty remains.",
                "The central question is not only whether {subject_l} is impressive, but whether it can reliably improve outcomes in real settings. Reliability, cost, latency, privacy, and maintainability determine whether a promising idea becomes useful infrastructure.",
                "Any serious discussion should also identify dependencies. Data quality, evaluation methods, security controls, and user workflows often matter as much as the model or tool itself. Without those foundations, technical progress can produce fragile results.",
                "The useful path forward is iterative: define the job, measure baseline performance, introduce the capability carefully, and compare results against clear criteria. That approach keeps the analysis practical instead of speculative.",
            ],
            [
                "The core concepts behind {subject_l} include scope, autonomy, feedback, and integration. Scope defines what the system is responsible for; autonomy defines how much it can decide without approval; feedback determines how it learns from mistakes; integration determines whether it fits existing work.",
                "A common failure is treating capability as a single number. In practice, a system can be strong in one environment and weak in another. The surrounding process, available context, and tolerance for error change the meaning of performance.",
                "Evaluation therefore has to be task-specific. Benchmarks are useful, but they are not substitutes for measuring the work people actually need done. The closer the evaluation is to the real workflow, the more trustworthy the result becomes.",
                "These concepts help teams avoid hype and panic. They create a vocabulary for deciding what to build, what to automate, what to monitor, and what should remain under human judgment.",
            ],
            [
                "The current state of {subject_l} is best described as uneven but significant. Some capabilities are mature enough for everyday use, while others remain experimental, brittle, or expensive. That mixed reality calls for careful adoption rather than blanket enthusiasm or rejection.",
                "Organizations benefit when they start with narrow use cases where success is easy to observe. Drafting, summarization, classification, retrieval, monitoring, and structured assistance can often show value before broader automation is attempted.",
                "The risks are equally concrete. Incorrect outputs, hidden assumptions, data leakage, dependency on vendors, and weak observability can turn a useful tool into an operational liability. These risks can be managed, but only if they are named early.",
                "A balanced current-state assessment should therefore combine ambition with discipline. The technology may open new possibilities, but durable value comes from design, governance, and verification.",
            ],
        ]
        expansion = [
            "For {heading_l}, the important step is to connect the concept to an observable workflow. A capability is only useful when it changes quality, speed, cost, reliability, or access in a way that can be measured. That measurement should be simple enough to repeat and specific enough to guide decisions.",
            "Governance also matters. Teams need clear ownership, escalation paths, logging, and review points so that {subject_l} does not become a black box. The stronger the operational discipline, the easier it becomes to experiment without losing control of risk.",
            "A phased approach is usually best. Start with low-risk tasks, compare outputs against a baseline, collect failures, and widen the scope only when performance is stable. This keeps the work grounded in evidence rather than excitement.",
            "The long-term value of {subject_l} depends on trust. Trust is built when users understand what the system can do, where it can fail, and how those failures will be caught. Technical quality and human confidence have to grow together.",
        ]
    else:
        headings = [
            "Overview",
            "Background",
            "Key Themes",
            "Practical Implications",
            "Challenges",
            "Examples",
            "Recommendations",
            "Conclusion",
        ]
        paragraph_sets = [
            [
                "{subject} can be understood by looking at its context, consequences, and practical meaning. A useful document should move beyond generic description and explain why the topic matters, what tradeoffs it creates, and how people can think about it with clarity.",
                "The background matters because every topic arrives with assumptions. Those assumptions influence what people notice, what they ignore, and what they consider a good result. Naming them makes the discussion more useful.",
                "A practical view also asks who is affected. Different people may experience the same subject as an opportunity, a risk, a responsibility, or a source of uncertainty. Good analysis keeps those perspectives visible.",
                "The result should be specific enough to support action. Clear definitions, examples, constraints, and next steps turn a broad topic into something that can be discussed and improved.",
            ],
            [
                "The key themes surrounding {subject_l} include purpose, impact, complexity, and accountability. Purpose explains why the topic matters; impact describes what changes; complexity identifies what makes the subject hard; accountability names who must respond.",
                "Examples are important because they make abstract points testable. A claim that sounds persuasive in general may look different when applied to a real person, team, process, or decision.",
                "Challenges should not be treated as reasons to stop thinking. They are signals that the topic needs better structure, clearer priorities, and more honest measurement.",
                "A strong conclusion should return to the original purpose. It should show what has been clarified, what remains uncertain, and what a sensible next step would look like.",
            ],
        ]
        expansion = [
            "In relation to {heading_l}, the topic becomes clearer when it is connected to concrete examples. Examples show where an idea succeeds, where it breaks down, and what details are easy to miss in a broad summary.",
            "Another useful lens is consequence. {subject} may affect decisions, relationships, resources, or expectations in ways that are not immediately visible. Exploring those consequences helps turn a general topic into a practical analysis.",
            "The topic also benefits from comparison. Looking at alternatives, tradeoffs, and edge cases prevents the document from becoming one-sided. A balanced account explains both the appeal of an idea and the reasons someone might approach it carefully.",
            "The strongest conclusion is one that returns to action. After the context and tradeoffs are clear, the reader should understand what matters most, what remains uncertain, and what reasonable next step follows from the analysis.",
        ]

    while len(headings) < page_count:
        headings.append(f"Extended Perspective {len(headings) + 1}")

    section_count = max(1, min(page_count, len(headings)))
    section_target_words = max(430, target_word_count(page_count) // max(1, page_count))
    sections: list[Section] = []
    for index, heading in enumerate(headings[:section_count]):
        source = paragraph_sets[index % len(paragraph_sets)]
        paragraphs = [
            paragraph.format(subject=subject_display, subject_l=subject_l)
            for paragraph in source
        ]
        while sum(len(paragraph.split()) for paragraph in paragraphs) < section_target_words:
            template = expansion[(len(paragraphs) + index) % len(expansion)]
            paragraphs.append(
                template.format(
                    subject=subject_display,
                    subject_l=subject_l,
                    heading_l=heading.lower(),
                )
            )
        if heading.lower().startswith("conclusion"):
            paragraphs.append(
                (
                    "{subject} ultimately matters because it gathers feeling, thought, and action "
                    "into one human problem: how to live with care. The answer is never complete, "
                    "but the attempt to answer it shapes character, relationships, and the quality "
                    "of everyday life."
                ).format(subject=subject_display)
            )
        sections.append(Section(heading=heading, paragraphs=paragraphs))
    return GeneratedDocument(title=title, sections=sections)


def document_word_count(doc: GeneratedDocument) -> int:
    total = len(doc.title.split())
    for section in doc.sections:
        total += len(section.heading.split())
        total += sum(len(paragraph.split()) for paragraph in section.paragraphs)
    return total


def generate_document(goal: str, page_count: int) -> GeneratedDocument:
    fallback_title = title_from_goal(goal)
    api_key = minimax_key_from_config()
    if api_key:
        try:
            raw = call_minimax(goal, page_count, api_key)
            parsed = extract_json_object(raw)
            if parsed:
                generated = document_from_json(parsed, fallback_title)
                if generated:
                    return generated
            return document_from_plain_text(raw, fallback_title)
        except Exception as exc:  # noqa: BLE001
            print(f"warning=AI content generation failed, using local fallback: {exc}", file=sys.stderr)
    return fallback_document(goal, page_count)


def ensure_document_shape(doc: GeneratedDocument, goal: str, page_count: int) -> GeneratedDocument:
    filler = fallback_document(goal, page_count)
    target_words = target_word_count(page_count)

    if len(doc.sections) < max(1, page_count // 3) or document_word_count(doc) < target_words * 0.5:
        return filler

    existing = {section.heading.lower() for section in doc.sections}
    for section in filler.sections:
        if len(doc.sections) >= page_count:
            break
        if section.heading.lower() not in existing:
            doc.sections.append(section)
            existing.add(section.heading.lower())

    for index, section in enumerate(doc.sections[:page_count]):
        if sum(len(paragraph.split()) for paragraph in section.paragraphs) >= 220:
            continue
        replacement = filler.sections[index % len(filler.sections)]
        section.paragraphs.extend(replacement.paragraphs)

    while len(doc.sections) < page_count:
        idx = len(doc.sections) + 1
        doc.sections.append(filler.sections[(idx - 1) % len(filler.sections)])
    return doc


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Aptos Display"
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)


def write_docx(generated: GeneratedDocument, path: Path, page_count: int) -> None:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(generated.title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)

    document.add_paragraph()
    sections = generated.sections[: max(1, page_count)]
    for index, section in enumerate(sections):
        heading = document.add_heading(section.heading, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for paragraph in section.paragraphs:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(9)
            p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(paragraph)
        if index < min(page_count, len(sections)) - 1:
            document.add_page_break()

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--goal", required=True)
    parser.add_argument("--output-dir", default="")
    parser.add_argument("--open-when-done", action="store_true")
    args = parser.parse_args(argv)

    goal = normalize_goal(args.goal)
    page_count = page_count_from_goal(args.goal)
    generated = ensure_document_shape(generate_document(goal, page_count), goal, page_count)
    out_dir = output_dir(args.output_dir or None)
    out_path = out_dir / f"{safe_filename(generated.title)}.docx"
    write_docx(generated, out_path, page_count)

    if args.open_when_done:
        try:
            os.startfile(out_path)  # type: ignore[attr-defined]
        except OSError as exc:
            print(f"warning=Document created but auto-open failed: {exc}", file=sys.stderr)

    print(f"created_document={out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
